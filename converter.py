import pdfplumber
from docx import Document
import sys
import os
import io
import re

def is_in_bbox(element_bbox, container_bbox):
    """Checks if element_bbox is inside container_bbox."""
    x0, y0, x1, y1 = element_bbox
    cx0, cy0, cx1, cy1 = container_bbox
    return x0 >= cx0 and y0 >= cy0 and x1 <= cx1 and y1 <= cy1

def convert_pdf_to_docx(pdf_path, docx_path):
    """
    Converts a PDF file to a .docx file, with improved paragraph and list detection.
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            document = Document()
            document.add_heading(f"Conversion of {os.path.basename(pdf_path)}", level=1)

            for i, page in enumerate(pdf.pages):
                document.add_heading(f"--- Page {i+1} ---", level=2)

                elements = []
                table_bboxes = []

                # Extract tables
                found_tables = page.find_tables()
                if found_tables:
                    extracted_tables_content = page.extract_tables()
                    if len(found_tables) == len(extracted_tables_content):
                        for i_tbl, tbl_obj in enumerate(found_tables):
                            table_bboxes.append(tbl_obj.bbox)
                            elements.append({
                                'type': 'table', 'y0': tbl_obj.bbox[1], 'bbox': tbl_obj.bbox,
                                'data': extracted_tables_content[i_tbl]
                            })

                # Extract images
                for img in page.images:
                    elements.append({
                        'type': 'image', 'y0': img['top'], 'bbox': (img['x0'], img['top'], img['x1'], img['bottom']),
                        'data': img
                    })

                # Extract words and form lines
                words = page.extract_words(x_tolerance=1, y_tolerance=1)
                lines = {}
                for word in words:
                    word_bbox = (word['x0'], word['top'], word['x1'], word['bottom'])
                    if any(is_in_bbox(word_bbox, tbl_bbox) for tbl_bbox in table_bboxes):
                        continue
                    line_top = round(word['top'])
                    if line_top not in lines: lines[line_top] = []
                    lines[line_top].append(word)

                sorted_lines = []
                for line_top in sorted(lines.keys()):
                    line_words = sorted(lines[line_top], key=lambda w: w['x0'])
                    sorted_lines.append({
                        'text': " ".join(w['text'] for w in line_words),
                        'bbox': (
                            min(w['x0'] for w in line_words), min(w['top'] for w in line_words),
                            max(w['x1'] for w in line_words), max(w['bottom'] for w in line_words)
                        )
                    })

                # Group lines into blocks (paragraphs or lists)
                line_blocks = []
                if sorted_lines:
                    current_block = [sorted_lines[0]]
                    for i_line in range(1, len(sorted_lines)):
                        line = sorted_lines[i_line]
                        prev_line = sorted_lines[i_line-1]
                        vertical_gap = line['bbox'][1] - prev_line['bbox'][3]
                        line_height = prev_line['bbox'][3] - prev_line['bbox'][1]
                        if vertical_gap < (line_height * 0.7):
                            current_block.append(line)
                        else:
                            line_blocks.append(current_block)
                            current_block = [line]
                    line_blocks.append(current_block)

                # Process blocks
                for block in line_blocks:
                    block_bbox = (
                        min(l['bbox'][0] for l in block), min(l['bbox'][1] for l in block),
                        max(l['bbox'][2] for l in block), max(l['bbox'][3] for l in block)
                    )

                    # List detection
                    is_bullet_list = all(re.match(r'^\s*[\*•-]\s+', l['text']) for l in block)
                    is_numbered_list = all(re.match(r'^\s*\d+\.\s+', l['text']) for l in block)

                    if is_bullet_list or is_numbered_list:
                        elements.append({
                            'type': 'list', 'y0': block_bbox[1], 'bbox': block_bbox,
                            'data': {
                                'items': [l['text'].strip() for l in block],
                                'style': 'List Bullet' if is_bullet_list else 'List Number'
                            }
                        })
                    else: # It's a paragraph
                        para_text = " ".join(l['text'] for l in block)
                        elements.append({
                            'type': 'text', 'y0': block_bbox[1], 'bbox': block_bbox, 'data': para_text
                        })

                # Sort all elements and add to document
                elements.sort(key=lambda el: el['y0'])
                for el in elements:
                    if el['type'] == 'text':
                        document.add_paragraph(el['data'])
                    elif el['type'] == 'list':
                        for item in el['data']['items']:
                            document.add_paragraph(item, style=el['data']['style'])
                    elif el['type'] == 'table':
                        # (Table adding logic as before)
                        table_data = el['data']
                        if not table_data: continue
                        num_rows, num_cols = len(table_data), len(table_data[0]) if table_data else 0
                        if num_cols == 0: continue
                        table = document.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'
                        for r_idx, r_data in enumerate(table_data):
                            if r_data:
                                for c_idx, c_text in enumerate(r_data):
                                    if c_idx < num_cols:
                                        table.cell(r_idx, c_idx).text = str(c_text) if c_text is not None else ''
                    elif el['type'] == 'image':
                        # (Image adding logic as before)
                        try:
                            img_data = el['data']
                            image_data = img_data.get("stream", {}).get("data", b"")
                            if not image_data: image_data = img_data.get('src')
                            if image_data:
                                document.add_picture(io.BytesIO(image_data))
                        except Exception as e:
                            print(f"Could not process image on page {i+1}: {e}")

            document.save(docx_path)
            print(f"Successfully converted {pdf_path} to {docx_path}")
            return True

    except Exception as e:
        print(f"An error occurred: {e}")
        return False

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python converter.py <input_pdf_path> <output_docx_path>")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]

    if not os.path.exists(input_pdf):
        print(f"Error: The file {input_pdf} does not exist.")
        sys.exit(1)

    convert_pdf_to_docx(input_pdf, output_docx)
